from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from PIL import Image
from reportlab.pdfgen.canvas import Canvas

from app.main import app
from app.modules.private_library import (
    LibraryIngestCommand,
    LibraryLifecycleCommand,
    LibraryOCRCommand,
    LibraryTextReviewCommand,
    OCRTextResult,
    OpenAIResourceOCRProvider,
    PrivateLibrary,
    PrivateLibraryError,
    QuestionCandidateUpdate,
)
from app.modules.question_bank import QuestionBank
from app.modules.private_library.exports import LibraryExportError, LibraryTextRenderer


def make_library(tmp_path: Path) -> PrivateLibrary:
    return PrivateLibrary(tmp_path / "library.sqlite3", tmp_path / "files")


def ingest_command(**updates) -> LibraryIngestCommand:
    values = {
        "title": "函数资料",
        "rights_basis": "private_teaching_only",
        "rights_statement": "本人确认该资料仅上传至私人空间用于日常教学。",
        "rights_acknowledged": True,
        "owner_id": "owner_teacher",
    }
    values.update(updates)
    return LibraryIngestCommand(**values)


def docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_heading("高中数学讲义", level=1)
    document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def pdf_bytes(text: str) -> bytes:
    output = BytesIO()
    canvas = Canvas(output)
    canvas.drawString(72, 760, text)
    canvas.save()
    return output.getvalue()


def png_bytes() -> bytes:
    output = BytesIO()
    Image.new("RGB", (80, 60), "white").save(output, format="PNG")
    return output.getvalue()


def image_docx_bytes() -> bytes:
    document = Document()
    document.add_picture(BytesIO(png_bytes()))
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_is_extracted_but_remains_private(tmp_path: Path) -> None:
    library = make_library(tmp_path)

    item = library.ingest(
        ingest_command(),
        filename="函数讲义.docx",
        content=docx_bytes("函数单调性是本节课的核心内容。"),
    )

    assert item.file_kind == "docx"
    assert item.extraction_status == "extracted"
    assert "函数单调性" in item.extracted_text
    assert item.visibility == "private"
    assert not item.public_search_allowed
    assert not item.model_training_allowed
    assert not item.adaptation_allowed
    assert item.text_review_status == "pending"


def test_pdf_text_is_extracted_and_download_path_is_confined(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    item = library.ingest(
        ingest_command(rights_basis="original"),
        filename="lesson.pdf",
        content=pdf_bytes("Function monotonicity and x^2"),
    )

    assert item.extraction_status == "extracted"
    assert "Function monotonicity" in item.extracted_text
    assert item.page_count == 1
    assert item.adaptation_allowed
    path, downloaded = library.file_for_download(item.library_item_id)
    assert path.parent == (tmp_path / "files").resolve()
    assert path.read_bytes().startswith(b"%PDF-")
    assert downloaded.original_filename == "lesson.pdf"


def test_corrected_math_text_exports_to_readable_docx_and_pdf(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    item = library.ingest(ingest_command(), filename="lesson.pdf", content=pdf_bytes("x^2"))
    reviewed = library.review(item.library_item_id, LibraryTextReviewCommand(corrected_text=r"函数 $f(x)=\frac{x^2}{2}$", confirm=True))
    renderer = LibraryTextRenderer(tmp_path / "output")
    docx = renderer.render(reviewed, "docx")
    pdf = renderer.render(reviewed, "pdf")
    assert docx.path.read_bytes().startswith(b"PK")
    assert pdf.path.read_bytes().startswith(b"%PDF")
    document = Document(docx.path)
    assert "x²/2" in "\n".join(p.text for p in document.paragraphs)


def test_unreadable_font_text_cannot_be_exported_before_ocr(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    unreadable = library.ingest(
        ingest_command(rights_basis="original"),
        filename="scan.png",
        content=png_bytes(),
    )

    with pytest.raises(LibraryExportError, match="OCR"):
        LibraryTextRenderer(tmp_path / "output").render(unreadable, "docx")


def test_local_ocr_audit_does_not_claim_external_consent(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    item = library.ingest(ingest_command(), filename="lesson.pdf", content=pdf_bytes("x^2"))

    class FakeLocalProvider:
        name = "local_math_ocr"
        def extract(self, **_kwargs):
            return OCRTextResult(text=r"函数 $f(x)=x^2$", warnings=[])

    library.apply_local_ocr(item.library_item_id, provider=FakeLocalProvider(), teacher_id="teacher")
    import sqlite3
    with sqlite3.connect(tmp_path / "library.sqlite3") as connection:
        consent = connection.execute("SELECT external_consent FROM library_ocr_runs").fetchone()[0]
    assert consent == 0


def test_local_ocr_rejects_duplicate_long_running_job() -> None:
    from app.routes.library import _local_math_ocr_lock

    assert _local_math_ocr_lock.acquire(blocking=False)
    try:
        response = TestClient(app).post("/api/v1/library/any-item/local-math-ocr")
    finally:
        _local_math_ocr_lock.release()

    assert response.status_code == 409
    assert "正在处理" in response.json()["detail"]


def test_unreadable_pdf_font_text_cannot_be_falsely_confirmed(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    item = library.ingest(
        ingest_command(rights_basis="original"),
        filename="formula.pdf",
        content=pdf_bytes("Readable source page"),
    )

    with pytest.raises(PrivateLibraryError, match="不能确认"):
        library.review(
            item.library_item_id,
            LibraryTextReviewCommand(
                corrected_text="已知函数 f(\uf021)=\uf022，求 \uf023。", confirm=True
            ),
        )


def test_library_items_move_to_recycle_bin_and_restore(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    item = library.ingest(
        ingest_command(), filename="lesson.docx", content=docx_bytes("函数单调性")
    )

    trashed = library.change_lifecycle(
        item.library_item_id,
        LibraryLifecycleCommand(action="trash", reason="不再需要"),
    )
    assert trashed.lifecycle_state == "trashed"
    assert library.list().total == 0
    assert library.list(lifecycle_state="trashed").total == 1
    assert library.stats().trashed == 1

    restored = library.change_lifecycle(
        item.library_item_id, LibraryLifecycleCommand(action="restore")
    )
    assert restored.lifecycle_state == "active"
    assert library.list().total == 1


def test_image_waits_for_ocr_then_can_be_manually_confirmed(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    item = library.ingest(
        ingest_command(title="立体几何题图"),
        filename="geometry.png",
        content=png_bytes(),
    )

    assert item.file_kind == "image"
    assert item.extraction_status == "needs_ocr"
    assert item.extracted_text == ""
    assert any("人工转录" in warning for warning in item.warnings)

    reviewed = library.review(
        item.library_item_id,
        LibraryTextReviewCommand(
            corrected_text="如图，在正方体 ABCD-A₁B₁C₁D₁ 中，求异面直线所成角。",
            note="教师根据原图完成转录",
            confirm=True,
        ),
    )

    assert reviewed.text_review_status == "confirmed"
    assert reviewed.version == 2
    assert "正方体" in reviewed.corrected_text


def test_rights_acknowledgement_and_duplicate_gate(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    content = docx_bytes("集合与常用逻辑用语")

    with pytest.raises(PrivateLibraryError, match="必须确认"):
        library.ingest(
            ingest_command(rights_acknowledged=False),
            filename="sets.docx",
            content=content,
        )

    first = library.ingest(ingest_command(), filename="sets.docx", content=content)
    with pytest.raises(PrivateLibraryError, match=first.library_item_id):
        library.ingest(ingest_command(), filename="copy.docx", content=content)


def test_invalid_file_is_rejected_without_writing(tmp_path: Path) -> None:
    library = make_library(tmp_path)

    with pytest.raises(PrivateLibraryError, match="不支持的文件格式"):
        library.ingest(
            ingest_command(), filename="payload.exe", content=b"not-a-real-file"
        )

    assert library.stats().total == 0
    assert list((tmp_path / "files").iterdir()) == []


def test_private_library_http_upload_review_and_download(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    library = make_library(tmp_path)
    monkeypatch.setattr("app.routes.library.get_private_library", lambda: library)
    client = TestClient(app)

    response = client.post(
        "/api/v1/library",
        files={
            "file": (
                "函数讲义.docx",
                docx_bytes("函数的奇偶性与对称性。"),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "title": "函数奇偶性",
            "rights_basis": "private_teaching_only",
            "rights_statement": "本人确认该文件仅在私人空间用于教学备课。",
            "rights_acknowledged": "true",
        },
    )

    assert response.status_code == 201
    item = response.json()
    assert item["visibility"] == "private"
    assert "函数的奇偶性" in item["extracted_text"]

    review = client.patch(
        f"/api/v1/library/{item['library_item_id']}/review",
        json={
            "corrected_text": "函数的奇偶性可以借助图象对称性理解。",
            "note": "完成首轮校对",
            "confirm": True,
        },
    )
    download = client.get(f"/api/v1/library/{item['library_item_id']}/file")
    listing = client.get("/api/v1/library")

    assert review.status_code == 200
    assert review.json()["text_review_status"] == "confirmed"
    assert download.status_code == 200
    assert download.content.startswith(b"PK")
    assert listing.status_code == 200
    assert listing.json()["total"] == 1


def test_http_rejects_short_rights_statement(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    library = make_library(tmp_path)
    monkeypatch.setattr("app.routes.library.get_private_library", lambda: library)
    client = TestClient(app)

    response = client.post(
        "/api/v1/library",
        files={"file": ("lesson.docx", docx_bytes("函数"), "application/octet-stream")},
        data={
            "rights_basis": "private_teaching_only",
            "rights_statement": "太短",
            "rights_acknowledged": "true",
        },
    )

    assert response.status_code == 422
    assert library.stats().total == 0


def confirmed_question_resource(library: PrivateLibrary):
    text = """1. 已知集合 A={1,2}，则下列正确的是（ ）
A. 1∈A B. 3∈A C. A=∅ D. 2∉A
答案：A
解析：由集合元素可知 1∈A。
2. 已知函数 f(x)=x^2，求 f(2)。
答案：4
解析：代入 x=2，得 f(2)=4。"""
    item = library.ingest(
        ingest_command(title="集合与函数例题", rights_basis="original"),
        filename="questions.docx",
        content=docx_bytes(text),
    )
    return library.review(
        item.library_item_id,
        LibraryTextReviewCommand(corrected_text=text, confirm=True),
    )


def test_confirmed_text_becomes_editable_question_candidates(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    item = confirmed_question_resource(library)

    proposed = library.propose_questions(item.library_item_id)

    assert len(proposed.items) == 2
    assert proposed.items[0].question_type == "single_choice"
    assert [option.key for option in proposed.items[0].options] == ["A", "B", "C", "D"]
    assert proposed.items[0].answer_value == "A"
    assert proposed.items[1].answer_value == "4"

    updated = library.update_question_candidate(
        item.library_item_id,
        proposed.items[1].candidate_id,
        QuestionCandidateUpdate(**{
            **proposed.items[1].model_dump(),
            "stem_plain": "已知函数 f(x)=x²，求 f(3)。",
            "answer_value": "9",
            "final_answer": "9",
            "solution_steps": ["代入 x=3，得 f(3)=9。"],
        }),
    )
    assert updated.stem_plain.endswith("f(3)。")
    assert updated.answer_value == "9"


def test_candidate_import_is_private_auditable_and_idempotent(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    bank = QuestionBank(tmp_path / "questions.sqlite3", tmp_path / "media")
    item = confirmed_question_resource(library)
    candidate = library.propose_questions(item.library_item_id).items[0]

    question = bank.create_private_resource_question(
        candidate.model_dump(), resource=item.model_dump()
    )
    marked = library.mark_candidate_imported(
        item.library_item_id, candidate.candidate_id, question.question_id
    )
    same = bank.create_private_resource_question(candidate.model_dump(), resource=item.model_dump())

    assert question.visibility == "private"
    assert question.review_status == "pending"
    assert question.verification_status == "needs_math_review"
    assert "teacher_review_required" in question.publication_blockers
    assert question.raw["provenance"]["resource_candidate_id"] == candidate.candidate_id
    assert marked.status == "imported"
    assert same.question_id == question.question_id


def test_ocr_requires_per_run_consent_and_returns_pending_review(tmp_path: Path) -> None:
    library = make_library(tmp_path)
    item = library.ingest(
        ingest_command(title="扫描题目"), filename="scan.png", content=png_bytes()
    )

    class FakeOCR:
        name = "fake-ocr"

        def extract(self, **_kwargs):
            return OCRTextResult(text="1. 已知 a>0，求 a+a^{-1} 的最小值。", warnings=["第1页已识别"])

    with pytest.raises(PrivateLibraryError, match="明确同意"):
        library.apply_ocr(
            item.library_item_id,
            provider=FakeOCR(),
            consent=False,
            teacher_id="teacher",
        )

    reviewed, provider, warnings = library.apply_ocr(
        item.library_item_id,
        provider=FakeOCR(),
        consent=True,
        teacher_id="teacher",
    )
    assert provider == "fake-ocr"
    assert warnings == ["第1页已识别"]
    assert reviewed.extraction_status == "extracted"
    assert reviewed.text_review_status == "pending"
    assert "最小值" in reviewed.corrected_text


def test_openai_ocr_extracts_docx_embedded_images_as_vision_inputs() -> None:
    inputs = OpenAIResourceOCRProvider._media_inputs(
        image_docx_bytes(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="image-only.docx",
    )

    assert len(inputs) == 1
    assert inputs[0]["type"] == "input_image"
    assert inputs[0]["image_url"].startswith("data:image/png;base64,")


def test_question_candidate_http_flow(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    library = make_library(tmp_path)
    bank = QuestionBank(tmp_path / "questions.sqlite3", tmp_path / "media")
    item = confirmed_question_resource(library)
    monkeypatch.setattr("app.routes.library.get_private_library", lambda: library)
    monkeypatch.setattr("app.routes.library.get_library_question_bank", lambda: bank)
    client = TestClient(app)

    proposed = client.post(f"/api/v1/library/{item.library_item_id}/question-candidates")
    assert proposed.status_code == 200
    candidate = proposed.json()["items"][0]

    imported = client.post(
        f"/api/v1/library/{item.library_item_id}/question-candidates/{candidate['candidate_id']}/import"
    )
    repeated = client.post(
        f"/api/v1/library/{item.library_item_id}/question-candidates/{candidate['candidate_id']}/import"
    )
    assert imported.status_code == 200
    assert imported.json()["candidate"]["status"] == "imported"
    assert repeated.json()["already_imported"] is True


def test_image_resource_is_kept_as_a_bounded_question_image(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    library = make_library(tmp_path)
    bank = QuestionBank(tmp_path / "questions.sqlite3", tmp_path / "media")
    uploaded = library.ingest(
        ingest_command(title="立体几何原题图", rights_basis="original"),
        filename="solid.png",
        content=png_bytes(),
    )
    item = library.review(
        uploaded.library_item_id,
        LibraryTextReviewCommand(
            corrected_text="1. 如图，在正方体 ABCD-A₁B₁C₁D₁ 中，求异面直线所成角。\n答案：60°",
            confirm=True,
        ),
    )
    candidate = library.propose_questions(item.library_item_id).items[0]
    monkeypatch.setattr("app.routes.library.get_private_library", lambda: library)
    monkeypatch.setattr("app.routes.library.get_library_question_bank", lambda: bank)
    client = TestClient(app)

    response = client.post(
        f"/api/v1/library/{item.library_item_id}/question-candidates/{candidate.candidate_id}/import"
    )
    question = bank.get_question(response.json()["question_id"])

    assert response.status_code == 200
    assert len(question.images) == 1
    assert question.images[0].placement == "stem"
    assert question.images[0].width == 80
    assert question.images[0].height == 60
