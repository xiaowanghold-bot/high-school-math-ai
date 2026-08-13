from __future__ import annotations

import hashlib
import json
import re
import sqlite3
from zipfile import BadZipFile, ZipFile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from PIL import Image, UnidentifiedImageError
from pypdf import PdfReader

from app.modules.pdf_imports.text_repair import private_use_glyph_count

from app.modules.private_library.schemas import (
    LibraryIngestCommand,
    LibraryLifecycleCommand,
    QuestionCandidateList,
    QuestionCandidateOption,
    QuestionCandidateUpdate,
    QuestionCandidateView,
    LibraryItemList,
    LibraryItemSummary,
    LibraryItemView,
    LibraryStats,
    LibraryTextReviewCommand,
)


class PrivateLibraryError(ValueError):
    pass


class PrivateLibrary:
    """Owns private-file validation, extraction, rights and text review."""

    MAX_FILE_BYTES = 50 * 1024 * 1024
    MAX_TEXT_CHARS = 2_000_000
    MAX_PDF_PAGES = 1000
    MAX_DOCX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024
    MAX_IMAGE_PIXELS = 25_000_000
    MAX_CANDIDATES = 100

    def __init__(self, database_path: Path, file_root: Path) -> None:
        self.database_path = database_path.resolve()
        self.file_root = file_root.resolve()
        self.database_path.parent.mkdir(parents=True, exist_ok=True)
        self.file_root.mkdir(parents=True, exist_ok=True)
        self._initialize()

    def ingest(
        self,
        command: LibraryIngestCommand,
        *,
        filename: str,
        content: bytes,
    ) -> LibraryItemView:
        if not command.rights_acknowledged:
            raise PrivateLibraryError("必须确认资料来源与使用权声明")
        if not content:
            raise PrivateLibraryError("上传文件不能为空")
        if len(content) > self.MAX_FILE_BYTES:
            raise PrivateLibraryError("单个资料文件不能超过 50 MB")
        safe_filename = Path(filename or "未命名资料").name
        if safe_filename in {"", ".", ".."}:
            safe_filename = "未命名资料"
        extraction = self._extract(safe_filename, content)
        digest = hashlib.sha256(content).hexdigest()
        now = self._now()
        item_id = f"lib_{uuid4().hex[:16]}"
        title = command.title.strip() or Path(safe_filename).stem[:300] or "未命名资料"
        adaptation_allowed = command.rights_basis in {"original", "licensed"}
        with self._connect() as connection:
            duplicate = connection.execute(
                "SELECT library_item_id FROM library_items WHERE owner_id = ? AND source_sha256 = ?",
                (command.owner_id, digest),
            ).fetchone()
            if duplicate:
                raise PrivateLibraryError(
                    f"该文件已经上传，资料编号为 {duplicate['library_item_id']}"
                )
            stored_name = f"{item_id}{extraction['suffix']}"
            stored_path = self.file_root / stored_name
            stored_path.write_bytes(content)
            try:
                connection.execute(
                    """
                    INSERT INTO library_items (
                        library_item_id, owner_id, title, original_filename, stored_name,
                        file_kind, mime_type, size_bytes, source_sha256, page_count,
                        extraction_status, extracted_text, corrected_text, text_review_status,
                        rights_basis, rights_statement, adaptation_allowed, warnings_json,
                        review_note, version, created_at, updated_at
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, '', 'pending', ?, ?, ?, ?, '', 1, ?, ?)
                    """,
                    (
                        item_id,
                        command.owner_id,
                        title,
                        safe_filename,
                        stored_name,
                        extraction["file_kind"],
                        extraction["mime_type"],
                        len(content),
                        digest,
                        extraction["page_count"],
                        extraction["status"],
                        extraction["text"],
                        command.rights_basis,
                        command.rights_statement.strip(),
                        int(adaptation_allowed),
                        json.dumps(extraction["warnings"], ensure_ascii=False),
                        now,
                        now,
                    ),
                )
            except Exception:
                stored_path.unlink(missing_ok=True)
                raise
        return self.get(item_id)

    def list(self, *, limit: int = 50, lifecycle_state: str = "active") -> LibraryItemList:
        if lifecycle_state not in {"active", "trashed", "all"}:
            raise PrivateLibraryError("不支持的资料状态")
        where = "" if lifecycle_state == "all" else "WHERE lifecycle_state = ?"
        values = () if lifecycle_state == "all" else (lifecycle_state,)
        with self._connect() as connection:
            total = connection.execute(
                f"SELECT COUNT(*) FROM library_items {where}", values
            ).fetchone()[0]
            rows = connection.execute(
                f"SELECT * FROM library_items {where} ORDER BY updated_at DESC LIMIT ?",
                (*values, limit),
            ).fetchall()
        return LibraryItemList(items=[self._summary(row) for row in rows], total=total)

    def change_lifecycle(
        self, item_id: str, command: LibraryLifecycleCommand
    ) -> LibraryItemView:
        target = "trashed" if command.action == "trash" else "active"
        now = self._now()
        with self._connect() as connection:
            row = connection.execute(
                "SELECT lifecycle_state FROM library_items WHERE library_item_id = ?",
                (item_id,),
            ).fetchone()
            if row is None:
                raise KeyError(item_id)
            if row["lifecycle_state"] == target:
                return self.get(item_id)
            connection.execute(
                """
                UPDATE library_items
                SET lifecycle_state = ?, trashed_at = ?, lifecycle_reason = ?, updated_at = ?
                WHERE library_item_id = ?
                """,
                (
                    target,
                    now if target == "trashed" else None,
                    command.reason.strip(),
                    now,
                    item_id,
                ),
            )
            connection.execute(
                """
                INSERT INTO library_lifecycle_events
                (library_item_id, action, actor_id, reason, created_at)
                VALUES (?, ?, ?, ?, ?)
                """,
                (item_id, command.action, command.actor_id, command.reason.strip(), now),
            )
        return self.get(item_id)

    def get(self, item_id: str) -> LibraryItemView:
        with self._connect() as connection:
            row = connection.execute(
                "SELECT * FROM library_items WHERE library_item_id = ?", (item_id,)
            ).fetchone()
        if row is None:
            raise KeyError(item_id)
        return self._view(row)

    def review(self, item_id: str, command: LibraryTextReviewCommand) -> LibraryItemView:
        corrected = command.corrected_text.strip()
        if command.confirm and not corrected:
            raise PrivateLibraryError("确认校对前必须提供可用文本")
        if command.confirm:
            issue_count = self._unreadable_glyph_count(corrected)
            if issue_count:
                raise PrivateLibraryError(
                    f"文本仍含 {issue_count} 个不可读公式字符，不能确认；请对照原 PDF 重建为 LaTeX，或先运行视觉 OCR"
                )
        now = self._now()
        with self._connect() as connection:
            row = connection.execute(
                "SELECT * FROM library_items WHERE library_item_id = ?", (item_id,)
            ).fetchone()
            if row is None:
                raise KeyError(item_id)
            new_version = row["version"] + 1
            status = "confirmed" if command.confirm else "pending"
            connection.execute(
                """
                INSERT INTO library_text_revisions (
                    library_item_id, version, previous_text, revised_text,
                    review_status, reviewer_id, note, created_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    item_id,
                    new_version,
                    row["corrected_text"],
                    corrected,
                    status,
                    command.reviewer_id,
                    command.note.strip(),
                    now,
                ),
            )
            connection.execute(
                """
                UPDATE library_items SET corrected_text = ?, text_review_status = ?,
                    review_note = ?, version = ?, updated_at = ?
                WHERE library_item_id = ?
                """,
                (corrected, status, command.note.strip(), new_version, now, item_id),
            )
        return self.get(item_id)

    def apply_ocr(self, item_id: str, *, provider, consent: bool, teacher_id: str) -> tuple[LibraryItemView, str, list[str]]:
        """Run an explicitly authorized OCR provider and return a reviewable draft."""
        if not consent:
            raise PrivateLibraryError("必须明确同意本次将私人文件发送给已配置的 OCR 服务")
        return self._apply_ocr_provider(
            item_id, provider=provider, teacher_id=teacher_id, external_consent=True
        )

    def apply_local_ocr(self, item_id: str, *, provider, teacher_id: str) -> tuple[LibraryItemView, str, list[str]]:
        return self._apply_ocr_provider(
            item_id, provider=provider, teacher_id=teacher_id, external_consent=False
        )

    def _apply_ocr_provider(
        self, item_id: str, *, provider, teacher_id: str, external_consent: bool
    ) -> tuple[LibraryItemView, str, list[str]]:
        path, item = self.file_for_download(item_id)
        result = provider.extract(
            path=path,
            mime_type=item.mime_type,
            filename=item.original_filename,
            teacher_id=teacher_id,
        )
        text = result.text.strip()[: self.MAX_TEXT_CHARS]
        if not text:
            raise PrivateLibraryError("OCR 未返回可用文字，请改用人工转录")
        now = self._now()
        warning = "OCR 文字尚未经过教师核对，不会自动进入题库。"
        with self._connect() as connection:
            row = connection.execute(
                "SELECT * FROM library_items WHERE library_item_id = ?", (item_id,)
            ).fetchone()
            if row is None:
                raise KeyError(item_id)
            warnings = list(dict.fromkeys([*json.loads(row["warnings_json"]), *result.warnings, warning]))
            next_version = row["version"] + 1
            connection.execute(
                """
                INSERT INTO library_text_revisions (
                    library_item_id, version, previous_text, revised_text,
                    review_status, reviewer_id, note, created_at
                ) VALUES (?, ?, ?, ?, 'pending', ?, ?, ?)
                """,
                (
                    item_id,
                    next_version,
                    row["corrected_text"],
                    text,
                    teacher_id,
                    f"通过 {provider.name} 生成 OCR 待校对文本",
                    now,
                ),
            )
            connection.execute(
                """
                UPDATE library_items SET extraction_status = 'extracted', extracted_text = ?,
                    corrected_text = ?, text_review_status = 'pending', warnings_json = ?,
                    review_note = ?, version = ?, updated_at = ? WHERE library_item_id = ?
                """,
                (text, text, json.dumps(warnings, ensure_ascii=False), warning, next_version, now, item_id),
            )
            connection.execute(
                """
                INSERT INTO library_ocr_runs
                (library_item_id, provider, teacher_id, external_consent, warning_json, created_at)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (item_id, provider.name, teacher_id, int(external_consent), json.dumps(result.warnings, ensure_ascii=False), now),
            )
        return self.get(item_id), provider.name, result.warnings

    def propose_questions(self, item_id: str) -> QuestionCandidateList:
        """Create deterministic, editable candidates from teacher-confirmed text."""
        with self._connect() as connection:
            item = connection.execute(
                "SELECT * FROM library_items WHERE library_item_id = ?", (item_id,)
            ).fetchone()
            if item is None:
                raise KeyError(item_id)
            if item["text_review_status"] != "confirmed":
                raise PrivateLibraryError("只有教师已确认的文本才能生成拆题候选")
            existing = connection.execute(
                """
                SELECT * FROM library_question_candidates
                WHERE library_item_id = ? AND source_version = ? ORDER BY position
                """,
                (item_id, item["version"]),
            ).fetchall()
            if existing:
                return self._candidate_list(item_id, item["version"], existing)
            parts = self._split_questions(item["corrected_text"])
            if len(parts) > self.MAX_CANDIDATES:
                parts = parts[: self.MAX_CANDIDATES]
            now = self._now()
            for position, part in enumerate(parts, start=1):
                parsed = self._parse_candidate(part)
                connection.execute(
                    """
                    INSERT INTO library_question_candidates (
                        candidate_id, library_item_id, source_version, position, question_type,
                        stem_plain, stem_latex, options_json, answer_value, solution_method,
                        solution_steps_json, final_answer, difficulty, status, warnings_json,
                        imported_question_id, updated_at
                    ) VALUES (?, ?, ?, ?, ?, ?, NULL, ?, ?, ?, ?, ?, 3, 'draft', ?, NULL, ?)
                    """,
                    (
                        f"cand_{uuid4().hex[:16]}", item_id, item["version"], position,
                        parsed["question_type"], parsed["stem"],
                        json.dumps(parsed["options"], ensure_ascii=False), parsed["answer"],
                        "资料解析待教师核对", json.dumps(parsed["solution_steps"], ensure_ascii=False),
                        parsed["answer"], json.dumps(parsed["warnings"], ensure_ascii=False), now,
                    ),
                )
            rows = connection.execute(
                "SELECT * FROM library_question_candidates WHERE library_item_id = ? AND source_version = ? ORDER BY position",
                (item_id, item["version"]),
            ).fetchall()
        return self._candidate_list(item_id, item["version"], rows)

    def list_question_candidates(self, item_id: str) -> QuestionCandidateList:
        item = self.get(item_id)
        with self._connect() as connection:
            rows = connection.execute(
                "SELECT * FROM library_question_candidates WHERE library_item_id = ? AND source_version = ? ORDER BY position",
                (item_id, item.version),
            ).fetchall()
        return self._candidate_list(item_id, item.version, rows)

    def update_question_candidate(
        self, item_id: str, candidate_id: str, command: QuestionCandidateUpdate
    ) -> QuestionCandidateView:
        now = self._now()
        with self._connect() as connection:
            row = self._candidate_row(connection, item_id, candidate_id)
            if row["status"] == "imported":
                raise PrivateLibraryError("已导入题库的候选不可覆盖，请在题库审核台继续修改")
            connection.execute(
                """
                UPDATE library_question_candidates SET question_type = ?, stem_plain = ?,
                    stem_latex = ?, options_json = ?, answer_value = ?, solution_method = ?,
                    solution_steps_json = ?, final_answer = ?, difficulty = ?, status = 'draft',
                    updated_at = ? WHERE candidate_id = ?
                """,
                (
                    command.question_type, command.stem_plain.strip(), command.stem_latex,
                    json.dumps([item.model_dump() for item in command.options], ensure_ascii=False),
                    command.answer_value, command.solution_method,
                    json.dumps(command.solution_steps, ensure_ascii=False), command.final_answer,
                    command.difficulty, now, candidate_id,
                ),
            )
            updated = self._candidate_row(connection, item_id, candidate_id)
        return self._candidate_view(updated)

    def discard_question_candidate(self, item_id: str, candidate_id: str) -> QuestionCandidateView:
        with self._connect() as connection:
            row = self._candidate_row(connection, item_id, candidate_id)
            if row["status"] == "imported":
                raise PrivateLibraryError("已导入题库的候选不能丢弃")
            connection.execute(
                "UPDATE library_question_candidates SET status = 'discarded', updated_at = ? WHERE candidate_id = ?",
                (self._now(), candidate_id),
            )
            updated = self._candidate_row(connection, item_id, candidate_id)
        return self._candidate_view(updated)

    def get_question_candidate(self, item_id: str, candidate_id: str) -> tuple[QuestionCandidateView, LibraryItemView]:
        item = self.get(item_id)
        with self._connect() as connection:
            row = self._candidate_row(connection, item_id, candidate_id)
        return self._candidate_view(row), item

    def mark_candidate_imported(self, item_id: str, candidate_id: str, question_id: str) -> QuestionCandidateView:
        with self._connect() as connection:
            row = self._candidate_row(connection, item_id, candidate_id)
            if row["imported_question_id"] and row["imported_question_id"] != question_id:
                raise PrivateLibraryError("该候选已关联其他题库题目")
            connection.execute(
                "UPDATE library_question_candidates SET status = 'imported', imported_question_id = ?, updated_at = ? WHERE candidate_id = ?",
                (question_id, self._now(), candidate_id),
            )
            updated = self._candidate_row(connection, item_id, candidate_id)
        return self._candidate_view(updated)

    def stats(self) -> LibraryStats:
        with self._connect() as connection:
            active_filter = "lifecycle_state = 'active'"
            total = connection.execute(
                f"SELECT COUNT(*) FROM library_items WHERE {active_filter}"
            ).fetchone()[0]
            pending = connection.execute(
                f"SELECT COUNT(*) FROM library_items WHERE {active_filter} AND text_review_status = 'pending'"
            ).fetchone()[0]
            confirmed = connection.execute(
                f"SELECT COUNT(*) FROM library_items WHERE {active_filter} AND text_review_status = 'confirmed'"
            ).fetchone()[0]
            needs_ocr = connection.execute(
                f"SELECT COUNT(*) FROM library_items WHERE {active_filter} AND extraction_status = 'needs_ocr'"
            ).fetchone()[0]
            trashed = connection.execute(
                "SELECT COUNT(*) FROM library_items WHERE lifecycle_state = 'trashed'"
            ).fetchone()[0]
            rows = connection.execute(
                f"SELECT file_kind, COUNT(*) AS count FROM library_items WHERE {active_filter} GROUP BY file_kind"
            ).fetchall()
        return LibraryStats(
            total=total,
            pending_review=pending,
            confirmed=confirmed,
            needs_ocr=needs_ocr,
            trashed=trashed,
            by_file_kind={row["file_kind"]: row["count"] for row in rows},
        )

    def file_for_download(self, item_id: str) -> tuple[Path, LibraryItemView]:
        item = self.get(item_id)
        with self._connect() as connection:
            row = connection.execute(
                "SELECT stored_name FROM library_items WHERE library_item_id = ?", (item_id,)
            ).fetchone()
        path = (self.file_root / row["stored_name"]).resolve()
        if path.parent != self.file_root or not path.is_file():
            raise PrivateLibraryError("资料原文件不存在或存储路径异常")
        return path, item

    def _extract(self, filename: str, content: bytes) -> dict:
        if content.startswith(b"%PDF-"):
            return self._extract_pdf(content)
        if content.startswith(b"PK\x03\x04"):
            return self._extract_docx(content)
        try:
            with Image.open(BytesIO(content)) as image:
                image.verify()
            with Image.open(BytesIO(content)) as image:
                image_format = (image.format or "").upper()
                mime = Image.MIME.get(image_format)
                suffix = {"PNG": ".png", "JPEG": ".jpg", "WEBP": ".webp"}.get(
                    image_format
                )
                if not mime or not suffix:
                    raise PrivateLibraryError("图片仅支持 PNG、JPEG 和 WebP")
                width, height = image.size
                if width * height > self.MAX_IMAGE_PIXELS:
                    raise PrivateLibraryError("图片像素总量不能超过 2500 万")
            return {
                "file_kind": "image",
                "mime_type": mime,
                "suffix": suffix,
                "page_count": 1,
                "status": "needs_ocr",
                "text": "",
                "warnings": [
                    f"图片尺寸为 {width} × {height}；当前等待 OCR 或教师人工转录。",
                    "图片内容未经识别，不会自动进入题库或模型训练。",
                ],
            }
        except Image.DecompressionBombError as exc:
            raise PrivateLibraryError("图片像素规模异常，已拒绝解析") from exc
        except (UnidentifiedImageError, OSError):
            raise PrivateLibraryError(
                f"不支持的文件格式：{Path(filename).suffix or '无法识别'}；仅支持 PDF、DOCX、PNG、JPEG、WebP"
            ) from None

    def _extract_pdf(self, content: bytes) -> dict:
        warnings: list[str] = []
        try:
            reader = PdfReader(BytesIO(content), strict=False)
            if reader.is_encrypted and reader.decrypt("") == 0:
                raise PrivateLibraryError("PDF 已加密，无法提取文本；请先解除密码后上传")
            if len(reader.pages) > self.MAX_PDF_PAGES:
                raise PrivateLibraryError(f"PDF 页数不能超过 {self.MAX_PDF_PAGES} 页")
            pages: list[str] = []
            extracted_chars = 0
            for index, page in enumerate(reader.pages, start=1):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    pages.append(f"【第 {index} 页】\n{page_text}")
                    extracted_chars += len(page_text)
                if extracted_chars > self.MAX_TEXT_CHARS:
                    warnings.append("提取文本超过 200 万字符，已截断等待人工分批处理。")
                    break
            text = "\n\n".join(pages)[: self.MAX_TEXT_CHARS]
            status = "extracted" if text.strip() else "needs_ocr"
            unreadable_count = self._unreadable_glyph_count(text)
            if unreadable_count:
                status = "needs_ocr"
                warnings.append(
                    f"检测到 {unreadable_count} 个 PDF 私有字体/缺失映射字符；自动文本不可直接使用。请以原页预览为准，通过视觉 OCR 或人工 LaTeX 重建后再确认。"
                )
            if status == "needs_ocr" and not text.strip():
                warnings.append("未检测到可复制文本，该 PDF 可能是扫描件，需要 OCR 或人工转录。")
            return {
                "file_kind": "pdf",
                "mime_type": "application/pdf",
                "suffix": ".pdf",
                "page_count": len(reader.pages),
                "status": status,
                "text": text,
                "warnings": warnings,
            }
        except PrivateLibraryError:
            raise
        except Exception as exc:
            raise PrivateLibraryError(f"PDF 文件损坏或无法读取：{exc}") from exc

    def _extract_docx(self, content: bytes) -> dict:
        try:
            with ZipFile(BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > 10_000:
                    raise PrivateLibraryError("DOCX 内部文件数量异常，已拒绝解析")
                if sum(entry.file_size for entry in entries) > self.MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise PrivateLibraryError("DOCX 解压后超过 200 MB，已拒绝解析")
                if "word/document.xml" not in archive.namelist():
                    raise PrivateLibraryError("上传的 ZIP 文件不是有效 DOCX 文档")
        except BadZipFile as exc:
            raise PrivateLibraryError("DOCX 压缩结构损坏") from exc
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise PrivateLibraryError(f"DOCX 文件损坏或无法读取：{exc}") from exc
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append("\t".join(cells))
        text = "\n".join(blocks)[: self.MAX_TEXT_CHARS]
        warnings: list[str] = []
        status = "extracted" if text.strip() else "needs_ocr"
        if status == "needs_ocr":
            warnings.append("DOCX 中未检测到正文，可能仅包含图片，需要 OCR 或人工转录。")
        if document.inline_shapes:
            warnings.append("文档包含图片；当前仅提取文字，图片中的题目仍需 OCR 或人工核对。")
        return {
            "file_kind": "docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "suffix": ".docx",
            "page_count": None,
            "status": status,
            "text": text,
            "warnings": warnings,
        }

    def _initialize(self) -> None:
        with self._connect() as connection:
            connection.executescript(
                """
                CREATE TABLE IF NOT EXISTS library_items (
                    library_item_id TEXT PRIMARY KEY,
                    owner_id TEXT NOT NULL,
                    title TEXT NOT NULL,
                    original_filename TEXT NOT NULL,
                    stored_name TEXT NOT NULL UNIQUE,
                    file_kind TEXT NOT NULL,
                    mime_type TEXT NOT NULL,
                    size_bytes INTEGER NOT NULL,
                    source_sha256 TEXT NOT NULL,
                    page_count INTEGER,
                    extraction_status TEXT NOT NULL,
                    extracted_text TEXT NOT NULL,
                    corrected_text TEXT NOT NULL,
                    text_review_status TEXT NOT NULL,
                    rights_basis TEXT NOT NULL,
                    rights_statement TEXT NOT NULL,
                    adaptation_allowed INTEGER NOT NULL,
                    warnings_json TEXT NOT NULL,
                    review_note TEXT NOT NULL,
                    version INTEGER NOT NULL,
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL,
                    lifecycle_state TEXT NOT NULL DEFAULT 'active',
                    trashed_at TEXT,
                    lifecycle_reason TEXT NOT NULL DEFAULT '',
                    UNIQUE(owner_id, source_sha256)
                );
                CREATE TABLE IF NOT EXISTS library_text_revisions (
                    revision_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    library_item_id TEXT NOT NULL,
                    version INTEGER NOT NULL,
                    previous_text TEXT NOT NULL,
                    revised_text TEXT NOT NULL,
                    review_status TEXT NOT NULL,
                    reviewer_id TEXT NOT NULL,
                    note TEXT NOT NULL,
                    created_at TEXT NOT NULL,
                    FOREIGN KEY(library_item_id) REFERENCES library_items(library_item_id),
                    UNIQUE(library_item_id, version)
                );
                CREATE TABLE IF NOT EXISTS library_ocr_runs (
                    run_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    library_item_id TEXT NOT NULL,
                    provider TEXT NOT NULL,
                    teacher_id TEXT NOT NULL,
                    external_consent INTEGER NOT NULL,
                    warning_json TEXT NOT NULL,
                    created_at TEXT NOT NULL,
                    FOREIGN KEY(library_item_id) REFERENCES library_items(library_item_id)
                );
                CREATE TABLE IF NOT EXISTS library_question_candidates (
                    candidate_id TEXT PRIMARY KEY,
                    library_item_id TEXT NOT NULL,
                    source_version INTEGER NOT NULL,
                    position INTEGER NOT NULL,
                    question_type TEXT NOT NULL,
                    stem_plain TEXT NOT NULL,
                    stem_latex TEXT,
                    options_json TEXT NOT NULL,
                    answer_value TEXT,
                    solution_method TEXT NOT NULL,
                    solution_steps_json TEXT NOT NULL,
                    final_answer TEXT,
                    difficulty INTEGER NOT NULL,
                    status TEXT NOT NULL,
                    warnings_json TEXT NOT NULL,
                    imported_question_id TEXT,
                    updated_at TEXT NOT NULL,
                    FOREIGN KEY(library_item_id) REFERENCES library_items(library_item_id),
                    UNIQUE(library_item_id, source_version, position)
                );
                CREATE INDEX IF NOT EXISTS idx_library_candidates_item
                    ON library_question_candidates(library_item_id, source_version, position);
                CREATE TABLE IF NOT EXISTS library_lifecycle_events (
                    event_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    library_item_id TEXT NOT NULL,
                    action TEXT NOT NULL,
                    actor_id TEXT NOT NULL,
                    reason TEXT NOT NULL,
                    created_at TEXT NOT NULL,
                    FOREIGN KEY(library_item_id) REFERENCES library_items(library_item_id)
                );
                """
            )
            columns = {
                row["name"] for row in connection.execute("PRAGMA table_info(library_items)")
            }
            if "lifecycle_state" not in columns:
                connection.execute(
                    "ALTER TABLE library_items ADD COLUMN lifecycle_state TEXT NOT NULL DEFAULT 'active'"
                )
            if "trashed_at" not in columns:
                connection.execute("ALTER TABLE library_items ADD COLUMN trashed_at TEXT")
            if "lifecycle_reason" not in columns:
                connection.execute(
                    "ALTER TABLE library_items ADD COLUMN lifecycle_reason TEXT NOT NULL DEFAULT ''"
                )
            rows = connection.execute(
                "SELECT library_item_id, extracted_text, corrected_text, warnings_json FROM library_items"
            ).fetchall()
            for row in rows:
                unreadable_count = self._unreadable_glyph_count(
                    row["corrected_text"] or row["extracted_text"]
                )
                if not unreadable_count:
                    continue
                warning = (
                    f"检测到 {unreadable_count} 个 PDF 私有字体/缺失映射字符；"
                    "已撤回错误的可用确认，请以原页预览为准并重建公式。"
                )
                warnings = list(dict.fromkeys([*json.loads(row["warnings_json"]), warning]))
                connection.execute(
                    """
                    UPDATE library_items SET extraction_status = 'needs_ocr',
                        text_review_status = 'pending', warnings_json = ?
                    WHERE library_item_id = ?
                    """,
                    (json.dumps(warnings, ensure_ascii=False), row["library_item_id"]),
                )

    def _connect(self) -> sqlite3.Connection:
        connection = sqlite3.connect(self.database_path)
        connection.row_factory = sqlite3.Row
        connection.execute("PRAGMA foreign_keys = ON")
        return connection

    @staticmethod
    def _now() -> str:
        return datetime.now(timezone.utc).isoformat()

    @staticmethod
    def _summary(row: sqlite3.Row) -> LibraryItemSummary:
        return LibraryItemSummary(
            library_item_id=row["library_item_id"],
            title=row["title"],
            original_filename=row["original_filename"],
            file_kind=row["file_kind"],
            mime_type=row["mime_type"],
            size_bytes=row["size_bytes"],
            page_count=row["page_count"],
            extraction_status=row["extraction_status"],
            text_review_status=row["text_review_status"],
            extracted_char_count=len(row["extracted_text"]),
            corrected_char_count=len(row["corrected_text"]),
            rights_basis=row["rights_basis"],
            lifecycle_state=row["lifecycle_state"],
            trashed_at=row["trashed_at"],
            version=row["version"],
            created_at=row["created_at"],
            updated_at=row["updated_at"],
        )

    @classmethod
    def _view(cls, row: sqlite3.Row) -> LibraryItemView:
        return LibraryItemView(
            **cls._summary(row).model_dump(),
            source_sha256=row["source_sha256"],
            extracted_text=row["extracted_text"],
            corrected_text=row["corrected_text"],
            rights_statement=row["rights_statement"],
            adaptation_allowed=bool(row["adaptation_allowed"]),
            warnings=json.loads(row["warnings_json"]),
            review_note=row["review_note"],
        )

    @staticmethod
    def _candidate_row(
        connection: sqlite3.Connection, item_id: str, candidate_id: str
    ) -> sqlite3.Row:
        row = connection.execute(
            "SELECT * FROM library_question_candidates WHERE library_item_id = ? AND candidate_id = ?",
            (item_id, candidate_id),
        ).fetchone()
        if row is None:
            raise KeyError(candidate_id)
        return row

    @classmethod
    def _candidate_list(
        cls, item_id: str, source_version: int, rows: list[sqlite3.Row]
    ) -> QuestionCandidateList:
        return QuestionCandidateList(
            library_item_id=item_id,
            source_version=source_version,
            items=[cls._candidate_view(row) for row in rows],
        )

    @staticmethod
    def _candidate_view(row: sqlite3.Row) -> QuestionCandidateView:
        return QuestionCandidateView(
            candidate_id=row["candidate_id"],
            library_item_id=row["library_item_id"],
            source_version=row["source_version"],
            position=row["position"],
            question_type=row["question_type"],
            stem_plain=row["stem_plain"],
            stem_latex=row["stem_latex"],
            options=[QuestionCandidateOption(**item) for item in json.loads(row["options_json"])],
            answer_value=row["answer_value"],
            solution_method=row["solution_method"],
            solution_steps=json.loads(row["solution_steps_json"]),
            final_answer=row["final_answer"],
            difficulty=row["difficulty"],
            status=row["status"],
            warnings=json.loads(row["warnings_json"]),
            imported_question_id=row["imported_question_id"],
            updated_at=row["updated_at"],
        )

    @staticmethod
    def _unreadable_glyph_count(text: str) -> int:
        return private_use_glyph_count(text) + len(
            re.findall(r"[�□■]|〔公式符号待核〕", text)
        )

    @staticmethod
    def _split_questions(text: str) -> list[str]:
        cleaned = re.sub(r"^【第\s*\d+\s*页】\s*$", "", text, flags=re.MULTILINE).strip()
        marker = re.compile(r"(?m)^\s*(?:第\s*)?\d{1,3}\s*[.、．]\s+")
        starts = [match.start() for match in marker.finditer(cleaned)]
        if len(starts) < 2:
            return [cleaned] if cleaned else []
        prefix = cleaned[: starts[0]].strip()
        parts: list[str] = []
        for index, start in enumerate(starts):
            end = starts[index + 1] if index + 1 < len(starts) else len(cleaned)
            part = cleaned[start:end].strip()
            if index == 0 and prefix and len(prefix) < 300:
                part = f"{prefix}\n{part}"
            if part:
                parts.append(part)
        return parts

    @staticmethod
    def _parse_candidate(text: str) -> dict:
        answer_match = re.search(r"(?ims)(?:参考)?答案\s*[：:]\s*(.+?)(?=\n\s*(?:解析|解答|解)\s*[：:]|\Z)", text)
        solution_match = re.search(r"(?ims)(?:解析|解答|解)\s*[：:]\s*(.+)\Z", text)
        answer = answer_match.group(1).strip() if answer_match else None
        solution = solution_match.group(1).strip() if solution_match else ""
        stem_end = min(
            [match.start() for match in (answer_match, solution_match) if match] or [len(text)]
        )
        stem_and_options = text[:stem_end].strip()
        option_pattern = re.compile(r"(?ms)(?:^|\s)([A-H])[.、．]\s*(.+?)(?=(?:\s+[A-H][.、．]\s)|\Z)")
        option_matches = list(option_pattern.finditer(stem_and_options))
        options = [
            {"key": match.group(1), "text": re.sub(r"\s+", " ", match.group(2)).strip()}
            for match in option_matches
        ]
        if option_matches:
            stem = stem_and_options[: option_matches[0].start()].strip()
        else:
            stem = stem_and_options
        stem = re.sub(r"(?m)^\s*(?:第\s*)?\d{1,3}\s*[.、．]\s*", "", stem, count=1).strip()
        warnings = ["自动拆题结果必须逐题核对题干、答案和解析。"]
        if not answer:
            warnings.append("未自动识别到答案，请教师补充。")
        if not solution:
            warnings.append("未自动识别到解析，请教师补充。")
        return {
            "question_type": "single_choice" if options else "open_response",
            "stem": stem or text.strip(),
            "options": options,
            "answer": answer,
            "solution_steps": [solution] if solution else [],
            "warnings": warnings,
        }
