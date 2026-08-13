from __future__ import annotations

import html
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from app.modules.exam_exports.math_text import teacher_readable_math
from app.modules.private_library.schemas import LibraryItemView


class LibraryExportError(RuntimeError):
    pass


@dataclass(frozen=True)
class RenderedLibraryText:
    path: Path
    media_type: str
    download_name: str


class LibraryTextRenderer:
    def __init__(self, output_root: Path) -> None:
        self.output_root = output_root

    def render(self, item: LibraryItemView, export_format: Literal["docx", "pdf"]) -> RenderedLibraryText:
        if item.extraction_status == "needs_ocr":
            raise LibraryExportError("当前文本包含 PDF 私有字体乱码，请先运行本地数学 OCR 或完成人工转录后再导出")
        text = item.corrected_text.strip()
        if not text:
            raise LibraryExportError("尚无可导出的校对文本，请先运行 OCR 或保存校对稿")
        target = self.output_root / "library" / export_format
        target.mkdir(parents=True, exist_ok=True)
        name = f"library-{item.library_item_id}-v{item.version}.{export_format}"
        path = target / name
        if export_format == "docx":
            self._docx(item, text, path)
            media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif export_format == "pdf":
            self._pdf(item, text, path)
            media = "application/pdf"
        else:
            raise LibraryExportError("不支持的导出格式")
        return RenderedLibraryText(path=path, media_type=media, download_name=name)

    @staticmethod
    def _docx(item: LibraryItemView, text: str, path: Path) -> None:
        document = Document()
        document.styles["Normal"].font.name = "Microsoft YaHei"
        document.styles["Normal"].font.size = Pt(11)
        document.add_heading(item.title, level=0)
        document.add_paragraph(f"教师校对版本 v{item.version} · 来源文件：{item.original_filename}")
        for line in text.split("\n"):
            document.add_paragraph(teacher_readable_math(line))
        document.save(path)

    @staticmethod
    def _font() -> Path:
        for path in (Path("C:/Windows/Fonts/msyh.ttc"), Path("C:/Windows/Fonts/simsun.ttc"), Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")):
            if path.is_file():
                return path
        raise LibraryExportError("未找到可用中文字体，请配置系统中文字体")

    def _pdf(self, item: LibraryItemView, text: str, path: Path) -> None:
        font_name = "LibraryCJK"
        if font_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(font_name, str(self._font()), subfontIndex=0))
        styles = getSampleStyleSheet()
        title = ParagraphStyle("LibraryTitle", parent=styles["Title"], fontName=font_name, fontSize=18, leading=26)
        meta = ParagraphStyle("LibraryMeta", parent=styles["BodyText"], fontName=font_name, fontSize=8, leading=13, textColor="#667085")
        body = ParagraphStyle("LibraryBody", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=17, spaceAfter=3)
        story = [Paragraph(html.escape(item.title), title), Paragraph(html.escape(f"教师校对版本 v{item.version} · 来源文件：{item.original_filename}"), meta), Spacer(1, 10)]
        for line in text.split("\n"):
            readable = teacher_readable_math(line)
            if readable.startswith("【第 ") and readable.endswith("页】") and len(story) > 3:
                story.append(PageBreak())
            story.append(Paragraph(html.escape(readable) or "&nbsp;", body))
        SimpleDocTemplate(str(path), pagesize=A4, leftMargin=42, rightMargin=42, topMargin=42, bottomMargin=42).build(story)
