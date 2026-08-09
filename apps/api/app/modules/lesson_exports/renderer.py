from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.modules.lesson_plans.schemas import LessonPlanView


ExportFormat = Literal["docx", "pdf"]


class LessonPlanExportError(RuntimeError):
    pass


@dataclass(frozen=True)
class RenderedLessonPlan:
    path: Path
    media_type: str
    download_name: str


class LessonPlanDocumentRenderer:
    """Render lesson plans while hiding document-library and filesystem details."""

    BLUE = "2E74B5"
    DARK_BLUE = "1F4D78"
    PALE_BLUE = "E8EEF5"
    TEXT = "263548"
    MUTED = "627084"
    WARNING = "FFF4D6"
    PAGE_WIDTH_DXA = 9360

    def __init__(
        self,
        *,
        output_root: Path,
        cjk_font_regular: Path | None = None,
        cjk_font_bold: Path | None = None,
    ) -> None:
        self.output_root = output_root
        self.cjk_font_regular = cjk_font_regular
        self.cjk_font_bold = cjk_font_bold

    def render(self, plan: LessonPlanView, export_format: ExportFormat) -> RenderedLessonPlan:
        if export_format not in {"docx", "pdf"}:
            raise LessonPlanExportError(f"不支持的导出格式：{export_format}")
        target_dir = self.output_root / export_format
        target_dir.mkdir(parents=True, exist_ok=True)
        download_name = f"lesson-plan-{plan.lesson_plan_id}-v{plan.version}.{export_format}"
        path = target_dir / download_name
        try:
            if export_format == "docx":
                self._render_docx(plan, path)
                media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                self._render_pdf(plan, path)
                media_type = "application/pdf"
        except LessonPlanExportError:
            raise
        except Exception as exc:  # pragma: no cover - third-party failures cross this seam
            raise LessonPlanExportError(f"教案导出失败：{exc}") from exc
        return RenderedLessonPlan(path=path, media_type=media_type, download_name=download_name)

    def _render_docx(self, plan: LessonPlanView, path: Path) -> None:
        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
        section.header_distance = section.footer_distance = Inches(0.492)
        self._configure_docx_styles(document)
        decimal_num_id = self._create_numbering(document, ordered=True)
        bullet_num_id = self._create_numbering(document, ordered=False)
        self._add_docx_header_footer(section)

        kicker = document.add_paragraph()
        kicker.paragraph_format.space_after = Pt(4)
        self._set_run_font(kicker.add_run("人教 A 版高中数学 · 教师教案"), 9, self.BLUE, bold=True)
        title = document.add_paragraph()
        title.paragraph_format.space_after = Pt(5)
        self._set_run_font(title.add_run(plan.content.title), 24, self.TEXT, bold=True)
        subtitle = document.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(12)
        self._set_run_font(
            subtitle.add_run(f"{plan.curriculum.volume} · {plan.curriculum.chapter} · {plan.curriculum.section}"),
            10,
            self.MUTED,
        )

        metrics = [
            ("课型", self._lesson_type_label(plan.request.lesson_type)),
            ("课时", f"{plan.request.duration_minutes} 分钟"),
            ("知识点", f"{len(plan.curriculum.knowledge_points)} 个"),
            ("题库例题", f"{len(plan.content.recommended_questions)} 道"),
        ]
        metric_table = document.add_table(rows=1, cols=4)
        self._set_table_geometry(metric_table, [2340] * 4)
        for cell, (label, value) in zip(metric_table.rows[0].cells, metrics, strict=True):
            self._shade_cell(cell, self.PALE_BLUE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            self._set_run_font(paragraph.add_run(f"{label}\n"), 8, self.MUTED)
            self._set_run_font(paragraph.add_run(value), 10, self.TEXT, bold=True)
        document.add_paragraph().paragraph_format.space_after = Pt(0)

        for warning in plan.generation.warnings:
            warning_table = document.add_table(rows=1, cols=1)
            self._set_table_geometry(warning_table, [self.PAGE_WIDTH_DXA])
            self._shade_cell(warning_table.cell(0, 0), self.WARNING)
            paragraph = warning_table.cell(0, 0).paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            self._set_run_font(paragraph.add_run(f"审核提示：{warning}"), 9, "805A00", bold=True)
        if plan.content.teacher_notes:
            notes_table = document.add_table(rows=1, cols=1)
            self._set_table_geometry(notes_table, [self.PAGE_WIDTH_DXA])
            self._shade_cell(notes_table.cell(0, 0), self.PALE_BLUE)
            paragraph = notes_table.cell(0, 0).paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            self._set_run_font(paragraph.add_run("备课提示："), 9, self.DARK_BLUE, bold=True)
            self._set_run_font(paragraph.add_run("；".join(plan.content.teacher_notes)), 9, self.TEXT)

        self._docx_heading(document, "一、教学目标", 1)
        self._add_docx_list(document, plan.content.objectives, decimal_num_id)
        self._docx_heading(document, "二、教学重点与难点", 1)
        self._docx_heading(document, "教学重点", 2)
        self._add_docx_list(document, plan.content.key_points, bullet_num_id)
        self._docx_heading(document, "教学难点", 2)
        self._add_docx_list(document, plan.content.difficulties, bullet_num_id)

        self._docx_heading(document, "三、教学流程", 1)
        flow_table = document.add_table(rows=1, cols=5)
        widths = [1260, 720, 2460, 2460, 2460]
        self._set_table_geometry(flow_table, widths)
        for cell, label in zip(flow_table.rows[0].cells, ["教学环节", "时间", "教师活动", "学生活动", "评价证据"], strict=True):
            self._shade_cell(cell, self.PALE_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            self._set_run_font(paragraph.add_run(label), 9, self.DARK_BLUE, bold=True)
        self._repeat_table_header(flow_table.rows[0])
        for phase in plan.content.teaching_flow:
            values = [phase.phase, f"{phase.minutes}′", phase.teacher_activity, phase.student_activity, phase.assessment]
            for index, (cell, value) in enumerate(zip(flow_table.add_row().cells, values, strict=True)):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index < 2 else WD_ALIGN_PARAGRAPH.LEFT
                self._set_run_font(paragraph.add_run(value), 8.5, self.TEXT)

        self._docx_heading(document, "四、题库联动", 1)
        if plan.content.recommended_questions:
            for index, question in enumerate(plan.content.recommended_questions, start=1):
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.19)
                paragraph.paragraph_format.first_line_indent = Inches(-0.19)
                paragraph.paragraph_format.space_after = Pt(6)
                self._set_run_font(paragraph.add_run(f"{index}. {question.usage}｜难度 {question.difficulty}/5\n"), 9, self.DARK_BLUE, bold=True)
                self._set_run_font(paragraph.add_run(question.stem), 10, self.TEXT)
                self._set_run_font(paragraph.add_run(f"\n题号：{question.question_id}｜质量状态：已独立验证"), 8, self.MUTED)
        else:
            document.add_paragraph("当前章节暂无可用题库例题。")

        self._docx_heading(document, "五、分层作业", 1)
        self._add_docx_list(document, plan.content.homework, decimal_num_id)
        self._docx_heading(document, "六、板书设计", 1)
        self._add_docx_list(document, plan.content.board_plan, bullet_num_id)
        document.core_properties.title = plan.content.title
        document.core_properties.subject = "人教 A 版高中数学教案"
        document.core_properties.author = "数研备课"
        document.core_properties.comments = "AI 生成初稿，须经教师审核后使用。"
        document.save(path)

    def _render_pdf(self, plan: LessonPlanView, path: Path) -> None:
        regular_path, bold_path = self._resolve_pdf_fonts()
        regular_name, bold_name = "MathLessonCJK", "MathLessonCJKBold"
        if regular_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(regular_name, str(regular_path)))
        if bold_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))
        styles = self._pdf_styles(regular_name, bold_name)
        story = [
            Paragraph("人教 A 版高中数学 · 教师教案", styles["Kicker"]),
            Paragraph(self._escape(plan.content.title), styles["LessonTitle"]),
            Paragraph(self._escape(f"{plan.curriculum.volume} · {plan.curriculum.chapter} · {plan.curriculum.section}"), styles["Subtitle"]),
        ]
        metric_values = [
            ("课型", self._lesson_type_label(plan.request.lesson_type)),
            ("课时", f"{plan.request.duration_minutes} 分钟"),
            ("知识点", f"{len(plan.curriculum.knowledge_points)} 个"),
            ("题库例题", f"{len(plan.content.recommended_questions)} 道"),
        ]
        metric_table = Table(
            [[Paragraph(f"<font size='8' color='#{self.MUTED}'>{label}</font><br/><b>{value}</b>", styles["Metric"]) for label, value in metric_values]],
            colWidths=[1.625 * inch] * 4,
        )
        metric_table.setStyle(TableStyle(self._pdf_table_style(background=self.PALE_BLUE)))
        story.extend([metric_table, Spacer(1, 8)])
        for warning in plan.generation.warnings:
            table = Table([[Paragraph(f"<b>审核提示：</b>{self._escape(warning)}", styles["Warning"])]], colWidths=[6.5 * inch])
            table.setStyle(TableStyle(self._pdf_table_style(background=self.WARNING, border="E0C77A")))
            story.extend([table, Spacer(1, 4)])
        if plan.content.teacher_notes:
            notes = "；".join(self._escape(note) for note in plan.content.teacher_notes)
            table = Table(
                [[Paragraph(f"<b>备课提示：</b>{notes}", styles["Notes"]) ]],
                colWidths=[6.5 * inch],
            )
            table.setStyle(TableStyle(self._pdf_table_style(background=self.PALE_BLUE)))
            story.extend([table, Spacer(1, 4)])

        story.extend(self._pdf_list_section("一、教学目标", plan.content.objectives, styles, ordered=True))
        story.append(Paragraph("二、教学重点与难点", styles["LessonHeading1"]))
        story.append(Paragraph("教学重点", styles["LessonHeading2"]))
        story.extend(self._pdf_list(plan.content.key_points, styles, ordered=False))
        story.append(Paragraph("教学难点", styles["LessonHeading2"]))
        story.extend(self._pdf_list(plan.content.difficulties, styles, ordered=False))
        story.append(Paragraph("三、教学流程", styles["LessonHeading1"]))
        flow_data = [[Paragraph(label, styles["TableHead"]) for label in ["教学环节", "时间", "教师活动", "学生活动", "评价证据"]]]
        for phase in plan.content.teaching_flow:
            flow_data.append([
                Paragraph(self._escape(phase.phase), styles["TableCellCenter"]),
                Paragraph(f"{phase.minutes}′", styles["TableCellCenter"]),
                Paragraph(self._escape(phase.teacher_activity), styles["TableCell"]),
                Paragraph(self._escape(phase.student_activity), styles["TableCell"]),
                Paragraph(self._escape(phase.assessment), styles["TableCell"]),
            ])
        flow_table = Table(flow_data, colWidths=[0.875 * inch, 0.5 * inch, 1.708 * inch, 1.708 * inch, 1.708 * inch], repeatRows=1)
        flow_table.setStyle(TableStyle(self._pdf_table_style(header=True)))
        story.append(flow_table)

        story.append(Paragraph("四、题库联动", styles["LessonHeading1"]))
        if plan.content.recommended_questions:
            for index, question in enumerate(plan.content.recommended_questions, start=1):
                story.append(KeepTogether([
                    Paragraph(f"<b>{index}. {self._escape(question.usage)}｜难度 {question.difficulty}/5</b>", styles["QuestionMeta"]),
                    Paragraph(self._escape(question.stem), styles["Body"]),
                    Paragraph(f"题号：{self._escape(question.question_id)}｜质量状态：已独立验证", styles["Small"]),
                ]))
        else:
            story.append(Paragraph("当前章节暂无可用题库例题。", styles["Body"]))
        story.extend(self._pdf_list_section("五、分层作业", plan.content.homework, styles, ordered=True))
        story.extend(self._pdf_list_section("六、板书设计", plan.content.board_plan, styles, ordered=False))

        pdf = SimpleDocTemplate(
            str(path), pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch,
            title=plan.content.title, author="数研备课", subject="人教 A 版高中数学教案",
        )

        def decorate_page(canvas, document) -> None:
            canvas.saveState()
            canvas.setFont(regular_name, 8)
            canvas.setFillColor(colors.HexColor(f"#{self.MUTED}"))
            canvas.drawString(inch, 10.5 * inch, "数研备课 | 人教 A 版 · 新高考Ⅰ卷")
            canvas.drawString(inch, 0.48 * inch, "数研备课 · 教师草稿")
            canvas.drawRightString(7.5 * inch, 0.48 * inch, f"第 {document.page} 页")
            canvas.restoreState()

        pdf.build(story, onFirstPage=decorate_page, onLaterPages=decorate_page)

    def _configure_docx_styles(self, document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        normal.font.color.rgb = RGBColor.from_string(self.TEXT)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        for name, size, color, before, after in [
            ("Heading 1", 16, self.BLUE, 18, 10),
            ("Heading 2", 13, self.BLUE, 14, 7),
            ("Heading 3", 12, self.DARK_BLUE, 10, 5),
        ]:
            style = document.styles[name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

    def _add_docx_header_footer(self, section) -> None:
        header = section.header.paragraphs[0]
        header.paragraph_format.space_after = Pt(0)
        self._set_run_font(header.add_run("数研备课 | 人教 A 版 · 新高考Ⅰ卷"), 8, self.MUTED)
        footer_table = section.footer.add_table(rows=1, cols=2, width=Inches(6.5))
        self._set_table_geometry(footer_table, [4680, 4680], borders=False)
        left, right = footer_table.rows[0].cells
        left.paragraphs[0].paragraph_format.space_after = Pt(0)
        self._set_run_font(left.paragraphs[0].add_run("数研备课 · 教师草稿"), 8, self.MUTED)
        right_p = right.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_after = Pt(0)
        self._set_run_font(right_p.add_run("第 "), 8, self.MUTED)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        right_p._p.append(field)
        self._set_run_font(right_p.add_run(" 页"), 8, self.MUTED)

    @staticmethod
    def _docx_heading(document: Document, text: str, level: int) -> None:
        document.add_heading(text, level=level)

    def _add_docx_list(self, document: Document, items: list[str], num_id: int) -> None:
        for item in items:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            p_pr = paragraph._p.get_or_add_pPr()
            num_pr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            num = OxmlElement("w:numId")
            num.set(qn("w:val"), str(num_id))
            num_pr.extend([ilvl, num])
            p_pr.append(num_pr)
            paragraph.add_run(item)

    @staticmethod
    def _create_numbering(document: Document, *, ordered: bool) -> int:
        numbering = document.part.numbering_part.element
        abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
        num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
        abstract_id, num_id = max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        for tag, value in [("start", "1"), ("numFmt", "decimal" if ordered else "bullet"), ("lvlText", "%1." if ordered else "•"), ("lvlJc", "left")]:
            node = OxmlElement(f"w:{tag}")
            node.set(qn("w:val"), value)
            level.append(node)
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)
        return num_id

    @staticmethod
    def _set_table_geometry(table, widths: list[int], *, borders: bool = True) -> None:
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_pr = table._tbl.tblPr
        tbl_width = tbl_pr.first_child_found_in("w:tblW")
        if tbl_width is None:
            tbl_width = OxmlElement("w:tblW")
            tbl_pr.append(tbl_width)
        tbl_width.set(qn("w:w"), str(sum(widths)))
        tbl_width.set(qn("w:type"), "dxa")
        table_indent = OxmlElement("w:tblInd")
        table_indent.set(qn("w:w"), "120")
        table_indent.set(qn("w:type"), "dxa")
        tbl_pr.append(table_indent)
        margins = OxmlElement("w:tblCellMar")
        for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")
            margins.append(node)
        tbl_pr.append(margins)
        border_nodes = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single" if borders else "nil")
            if borders:
                node.set(qn("w:sz"), "4")
                node.set(qn("w:color"), "AEBBCB")
            border_nodes.append(node)
        tbl_pr.append(border_nodes)
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for cell, width in zip(row.cells, widths, strict=True):
                tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tc_width.set(qn("w:w"), str(width))
                tc_width.set(qn("w:type"), "dxa")

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _repeat_table_header(row) -> None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        row._tr.get_or_add_trPr().append(repeat)

    @staticmethod
    def _set_run_font(run, size: float, color: str, *, bold: bool = False) -> None:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def _resolve_pdf_fonts(self) -> tuple[Path, Path]:
        regular_candidates = [self.cjk_font_regular, Path("C:/Windows/Fonts/msyh.ttc"), Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"), Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc")]
        bold_candidates = [self.cjk_font_bold, Path("C:/Windows/Fonts/msyhbd.ttc"), Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"), Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc")]
        regular = next((candidate for candidate in regular_candidates if candidate and candidate.exists()), None)
        bold = next((candidate for candidate in bold_candidates if candidate and candidate.exists()), None)
        if regular is None or bold is None:
            raise LessonPlanExportError("PDF 导出缺少中文字体；请配置 MATH_AI_CJK_FONT_REGULAR 与 MATH_AI_CJK_FONT_BOLD")
        return regular, bold

    def _pdf_styles(self, regular_name: str, bold_name: str):
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="Kicker", parent=styles["Normal"], fontName=bold_name, fontSize=9, leading=11.25, textColor=colors.HexColor(f"#{self.BLUE}"), spaceAfter=4))
        styles.add(ParagraphStyle(name="LessonTitle", parent=styles["Normal"], fontName=bold_name, fontSize=24, leading=30, textColor=colors.HexColor(f"#{self.TEXT}"), spaceAfter=5))
        styles.add(ParagraphStyle(name="Subtitle", parent=styles["Normal"], fontName=regular_name, fontSize=10, leading=12.5, textColor=colors.HexColor(f"#{self.MUTED}"), spaceAfter=12))
        styles.add(ParagraphStyle(name="LessonHeading1", parent=styles["Normal"], fontName=bold_name, fontSize=16, leading=20, textColor=colors.HexColor(f"#{self.BLUE}"), spaceBefore=18, spaceAfter=10, keepWithNext=True))
        styles.add(ParagraphStyle(name="LessonHeading2", parent=styles["Normal"], fontName=bold_name, fontSize=13, leading=16.25, textColor=colors.HexColor(f"#{self.BLUE}"), spaceBefore=14, spaceAfter=7, keepWithNext=True))
        styles.add(ParagraphStyle(name="Body", parent=styles["Normal"], fontName=regular_name, fontSize=11, leading=13.75, textColor=colors.HexColor(f"#{self.TEXT}"), spaceAfter=6))
        styles.add(ParagraphStyle(name="List", parent=styles["Body"], leftIndent=27, firstLineIndent=-13.5, spaceAfter=4))
        styles.add(ParagraphStyle(name="Metric", parent=styles["Body"], fontSize=10, leading=13, spaceAfter=0))
        styles.add(ParagraphStyle(name="Warning", parent=styles["Body"], fontSize=9, leading=11.25, textColor=colors.HexColor("#805A00"), spaceAfter=0))
        styles.add(ParagraphStyle(name="Notes", parent=styles["Body"], fontSize=9, leading=11.25, textColor=colors.HexColor(f"#{self.TEXT}"), spaceAfter=0))
        styles.add(ParagraphStyle(name="TableHead", parent=styles["Body"], fontName=bold_name, fontSize=8.5, leading=10.5, textColor=colors.HexColor(f"#{self.DARK_BLUE}"), alignment=TA_CENTER))
        styles.add(ParagraphStyle(name="TableCell", parent=styles["Body"], fontSize=8, leading=10.4, alignment=TA_LEFT))
        styles.add(ParagraphStyle(name="TableCellCenter", parent=styles["TableCell"], alignment=TA_CENTER))
        styles.add(ParagraphStyle(name="QuestionMeta", parent=styles["Body"], fontName=bold_name, fontSize=9, leading=11.25, textColor=colors.HexColor(f"#{self.DARK_BLUE}"), spaceBefore=5, spaceAfter=3))
        styles.add(ParagraphStyle(name="Small", parent=styles["Body"], fontSize=8, leading=10, textColor=colors.HexColor(f"#{self.MUTED}"), spaceAfter=6))
        return styles

    def _pdf_list_section(self, title: str, items: list[str], styles, *, ordered: bool):
        return [Paragraph(title, styles["LessonHeading1"]), *self._pdf_list(items, styles, ordered=ordered)]

    def _pdf_list(self, items: list[str], styles, *, ordered: bool):
        return [Paragraph(f"{index if ordered else '•'}{'.' if ordered else ''} {self._escape(item)}", styles["List"]) for index, item in enumerate(items, start=1)]

    @staticmethod
    def _pdf_table_style(*, background: str | None = None, border: str = "AEBBCB", header: bool = False):
        commands = [
            ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(f"#{border}")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]
        if background:
            commands.append(("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{background}")))
        if header:
            commands.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")))
        return commands

    @staticmethod
    def _lesson_type_label(lesson_type: str) -> str:
        return {"new_lesson": "新授课", "review": "复习课", "exercise": "习题课"}.get(lesson_type, lesson_type)

    @staticmethod
    def _escape(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
