from __future__ import annotations

import html
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.modules.exam_papers.schemas import (
    ExamPaperEdition,
    ExamPaperExportFormat,
    ExamPaperItemView,
    ExamPaperView,
)


class ExamPaperExportError(RuntimeError):
    pass


@dataclass(frozen=True)
class RenderedExamPaper:
    path: Path
    media_type: str
    download_name: str


class ExamPaperDocumentRenderer:
    """Renders student, answer and blueprint editions from immutable snapshots."""

    BLUE = "3559D9"
    TEXT = "172033"
    MUTED = "677287"
    LINE = "DDE3EC"
    PALE_BLUE = "EEF2FF"
    PALE_GREEN = "EEF8F4"

    def __init__(
        self,
        *,
        output_root: Path,
        asset_root: Path,
        cjk_font_regular: Path | None = None,
        cjk_font_bold: Path | None = None,
    ) -> None:
        self.output_root = output_root
        self.asset_root = asset_root.resolve()
        self.cjk_font_regular = cjk_font_regular
        self.cjk_font_bold = cjk_font_bold

    def render(
        self,
        paper: ExamPaperView,
        export_format: ExamPaperExportFormat,
        edition: ExamPaperEdition,
    ) -> RenderedExamPaper:
        if export_format not in {"docx", "pdf"}:
            raise ExamPaperExportError(f"不支持的导出格式：{export_format}")
        if edition not in {"student", "answer", "blueprint"}:
            raise ExamPaperExportError(f"不支持的试卷版本：{edition}")
        target_dir = self.output_root / "exam-papers" / export_format
        target_dir.mkdir(parents=True, exist_ok=True)
        download_name = (
            f"exam-paper-{paper.exam_paper_id}-v{paper.version}-{edition}.{export_format}"
        )
        path = target_dir / download_name
        try:
            if export_format == "docx":
                self._render_docx(paper, edition, path)
                media_type = (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                self._render_pdf(paper, edition, path)
                media_type = "application/pdf"
        except ExamPaperExportError:
            raise
        except Exception as exc:  # pragma: no cover - third-party failures cross the seam
            raise ExamPaperExportError(f"试卷导出失败：{exc}") from exc
        return RenderedExamPaper(path=path, media_type=media_type, download_name=download_name)

    def _render_docx(
        self, paper: ExamPaperView, edition: ExamPaperEdition, path: Path
    ) -> None:
        document = Document()
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = section.bottom_margin = Inches(0.72)
        section.left_margin = section.right_margin = Inches(0.78)
        self._configure_docx(document)
        self._docx_header_footer(section, edition)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_run(title.add_run(paper.title), 20, self.TEXT, bold=True)
        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_run(
            meta.add_run(
                f"考试时间：{paper.duration_minutes} 分钟　满分：{self._score(paper.total_score)} 分"
            ),
            10,
            self.MUTED,
        )
        if edition == "student":
            student = document.add_paragraph()
            student.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_run(student.add_run("班级：____________　姓名：____________　学号：____________"), 10, self.TEXT)
        notice = document.add_paragraph()
        notice.paragraph_format.space_after = Pt(12)
        self._set_run(notice.add_run(f"注意事项：{paper.instructions}"), 9, self.MUTED)

        if edition == "blueprint":
            self._docx_blueprint(document, paper)
        else:
            current_section = None
            for item in paper.items:
                if item.section_title != current_section:
                    current_section = item.section_title
                    heading = document.add_paragraph()
                    heading.paragraph_format.keep_with_next = True
                    self._set_run(heading.add_run(current_section), 14, self.BLUE, bold=True)
                self._docx_question(document, paper, item, edition)
            self._docx_sources(document, paper)

        document.core_properties.title = paper.title
        document.core_properties.subject = f"高中数学试卷（{self._edition_label(edition)}）"
        document.core_properties.author = "数研备课"
        document.core_properties.comments = "题目采用版本快照；AI 或规则生成内容须经教师审核。"
        document.save(path)

    def _docx_question(
        self,
        document: Document,
        paper: ExamPaperView,
        item: ExamPaperItemView,
        edition: ExamPaperEdition,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_with_next = bool(item.question.images)
        paragraph.paragraph_format.space_after = Pt(5)
        self._set_run(paragraph.add_run(f"{item.position}. "), 11, self.TEXT, bold=True)
        self._set_run(paragraph.add_run(item.question.stem_plain), 11, self.TEXT)
        self._set_run(paragraph.add_run(f"　（{self._score(item.score)} 分）"), 9, self.MUTED)
        for image in item.question.images:
            image_path = self._asset_path(paper.exam_paper_id, image.asset_id)
            width_inches = min(5.7, max(1.5, image.width / 150))
            height_inches = width_inches * image.height / image.width
            if height_inches > 3.6:
                width_inches *= 3.6 / height_inches
            picture = document.add_picture(str(image_path), width=Inches(width_inches))
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if image.caption:
                caption = document.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_run(caption.add_run(image.caption), 8, self.MUTED)
        for option in item.question.options:
            option_p = document.add_paragraph()
            option_p.paragraph_format.left_indent = Inches(0.28)
            option_p.paragraph_format.space_after = Pt(2)
            self._set_run(option_p.add_run(f"{option.key}. {option.text}"), 10.5, self.TEXT)
        if edition == "answer":
            answer = document.add_paragraph()
            answer.paragraph_format.left_indent = Inches(0.18)
            self._set_run(answer.add_run("参考答案："), 10, self.BLUE, bold=True)
            self._set_run(
                answer.add_run(item.question.final_answer or item.question.answer_value or "待教师补充"),
                10,
                self.TEXT,
            )
            if item.question.solution_method:
                method = document.add_paragraph()
                method.paragraph_format.left_indent = Inches(0.18)
                self._set_run(method.add_run(f"解析方法：{item.question.solution_method}"), 9, self.MUTED)
            for index, step in enumerate(item.question.solution_steps, start=1):
                step_p = document.add_paragraph()
                step_p.paragraph_format.left_indent = Inches(0.32)
                self._set_run(step_p.add_run(f"{index}. {step}"), 9.5, self.TEXT)
        elif item.question.question_type not in {"single_choice", "multiple_choice", "fill_blank"}:
            for _ in range(4):
                blank = document.add_paragraph("\n")
                blank.paragraph_format.space_after = Pt(2)

    def _docx_blueprint(self, document: Document, paper: ExamPaperView) -> None:
        heading = document.add_paragraph()
        self._set_run(heading.add_run("双向细目表"), 15, self.BLUE, bold=True)
        table = document.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        headers = ["题号", "题型", "章节", "知识点", "难度", "分值", "验证", "教师审核"]
        for cell, label in zip(table.rows[0].cells, headers, strict=True):
            self._shade_cell(cell, self.PALE_BLUE)
            self._set_run(cell.paragraphs[0].add_run(label), 8.5, self.BLUE, bold=True)
        for item in paper.items:
            values = [
                str(item.position),
                self._question_type_label(item.question.question_type),
                item.question.chapter or "未分类",
                "、".join(item.question.knowledge_point_ids) or "未标注",
                str(item.question.difficulty),
                self._score(item.score),
                "通过" if item.question.verification_status == "passed" else "待验算",
                "通过" if item.question.review_status == "approved" else "待审核",
            ]
            for cell, value in zip(table.add_row().cells, values, strict=True):
                self._set_run(cell.paragraphs[0].add_run(value), 8, self.TEXT)
        self._docx_breakdown(document, "章节结构", paper.chapter_breakdown)
        self._docx_breakdown(document, "难度结构", paper.difficulty_breakdown)
        self._docx_sources(document, paper)

    def _docx_breakdown(self, document: Document, title: str, rows) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        self._set_run(heading.add_run(title), 12, self.BLUE, bold=True)
        for row in rows:
            paragraph = document.add_paragraph()
            self._set_run(
                paragraph.add_run(
                    f"{row.label}：{row.question_count} 题，{self._score(row.score)} 分"
                ),
                9,
                self.TEXT,
            )

    def _docx_sources(self, document: Document, paper: ExamPaperView) -> None:
        document.add_page_break()
        heading = document.add_paragraph()
        self._set_run(heading.add_run("内容来源与审核说明"), 14, self.BLUE, bold=True)
        intro = document.add_paragraph()
        self._set_run(
            intro.add_run(
                "本试卷保存题目版本快照；题目来源仅用于教师审计，PDF 原版式与原解析未进入导出内容。"
            ),
            9,
            self.MUTED,
        )
        for item in paper.items:
            paragraph = document.add_paragraph()
            self._set_run(
                paragraph.add_run(
                    f"{item.position}. {item.question.question_id}｜{item.question.source_document}｜"
                    f"数学验证：通过｜教师审核：{'通过' if item.question.review_status == 'approved' else '待审核'}"
                ),
                8.5,
                self.TEXT,
            )

    def _render_pdf(
        self, paper: ExamPaperView, edition: ExamPaperEdition, path: Path
    ) -> None:
        regular_path, bold_path = self._resolve_pdf_fonts()
        regular_name, bold_name = "MathPaperCJK", "MathPaperCJKBold"
        if regular_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(regular_name, str(regular_path)))
        if bold_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))
        styles = self._pdf_styles(regular_name, bold_name)
        story = [
            Paragraph(self._escape(paper.title), styles["Title"]),
            Paragraph(
                f"考试时间：{paper.duration_minutes} 分钟　满分：{self._score(paper.total_score)} 分",
                styles["Meta"],
            ),
        ]
        if edition == "student":
            story.append(Paragraph("班级：____________　姓名：____________　学号：____________", styles["Meta"]))
        story.extend(
            [Paragraph(f"注意事项：{self._escape(paper.instructions)}", styles["Notice"]), Spacer(1, 8)]
        )
        if edition == "blueprint":
            story.extend(self._pdf_blueprint(paper, styles))
        else:
            current_section = None
            for item in paper.items:
                if item.section_title != current_section:
                    current_section = item.section_title
                    story.append(Paragraph(self._escape(current_section), styles["Heading"]))
                question_blocks = self._pdf_question(paper, item, edition, styles)
                if edition == "student" and item.question.question_type not in {
                    "single_choice",
                    "multiple_choice",
                    "fill_blank",
                }:
                    # Keep the response area with its prompt. A bare Spacer flowing to
                    # the next page otherwise produces a page that looks accidentally blank.
                    story.append(KeepTogether(question_blocks))
                else:
                    story.extend(question_blocks)
            story.extend([PageBreak(), *self._pdf_sources(paper, styles)])

        pdf = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=0.72 * inch,
            rightMargin=0.72 * inch,
            topMargin=0.72 * inch,
            bottomMargin=0.7 * inch,
            title=paper.title,
            author="数研备课",
            subject=f"高中数学试卷（{self._edition_label(edition)}）",
        )

        def decorate(canvas, document) -> None:
            canvas.saveState()
            canvas.setFont(regular_name, 8)
            canvas.setFillColor(colors.HexColor(f"#{self.MUTED}"))
            canvas.drawString(0.72 * inch, 0.38 * inch, f"数研备课 · {self._edition_label(edition)}")
            canvas.drawRightString(A4[0] - 0.72 * inch, 0.38 * inch, f"第 {document.page} 页")
            canvas.restoreState()

        pdf.build(story, onFirstPage=decorate, onLaterPages=decorate)

    def _pdf_question(self, paper, item, edition, styles) -> list:
        blocks = [
            Paragraph(
                f"<b>{item.position}.</b> {self._escape(item.question.stem_plain)}　"
                f"<font color='#{self.MUTED}'>（{self._score(item.score)} 分）</font>",
                styles["Question"],
            )
        ]
        for image in item.question.images:
            image_path = self._asset_path(paper.exam_paper_id, image.asset_id)
            width = min(5.6 * inch, max(1.5 * inch, image.width / 150 * inch))
            height = width * image.height / image.width
            if height > 3.5 * inch:
                width *= 3.5 * inch / height
                height = 3.5 * inch
            blocks.append(PdfImage(str(image_path), width=width, height=height))
            if image.caption:
                blocks.append(Paragraph(self._escape(image.caption), styles["Caption"]))
        for option in item.question.options:
            blocks.append(
                Paragraph(f"{self._escape(option.key)}. {self._escape(option.text)}", styles["Option"])
            )
        if edition == "answer":
            answer = self._escape(item.question.final_answer or item.question.answer_value or "待教师补充")
            blocks.append(Paragraph(f"<b>参考答案：</b>{answer}", styles["Answer"]))
            if item.question.solution_method:
                blocks.append(
                    Paragraph(f"解析方法：{self._escape(item.question.solution_method)}", styles["Small"])
                )
            for index, step in enumerate(item.question.solution_steps, start=1):
                blocks.append(Paragraph(f"{index}. {self._escape(step)}", styles["Solution"]))
        elif item.question.question_type not in {"single_choice", "multiple_choice", "fill_blank"}:
            blocks.append(Spacer(1, 1.25 * inch))
        blocks.append(Spacer(1, 7))
        return blocks

    def _pdf_blueprint(self, paper, styles) -> list:
        rows = [[Paragraph(label, styles["TableHead"]) for label in ["题号", "题型", "章节", "知识点", "难度", "分值", "验证", "审核"]]]
        for item in paper.items:
            values = [
                str(item.position),
                self._question_type_label(item.question.question_type),
                item.question.chapter or "未分类",
                "、".join(item.question.knowledge_point_ids) or "未标注",
                str(item.question.difficulty),
                self._score(item.score),
                "通过",
                "通过" if item.question.review_status == "approved" else "待审核",
            ]
            rows.append([Paragraph(self._escape(value), styles["TableCell"]) for value in values])
        table = Table(
            rows,
            colWidths=[0.35 * inch, 0.55 * inch, 1.15 * inch, 1.35 * inch, 0.35 * inch, 0.4 * inch, 0.45 * inch, 0.5 * inch],
            repeatRows=1,
        )
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{self.PALE_BLUE}")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor(f"#{self.LINE}")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        blocks = [Paragraph("双向细目表", styles["Heading"]), table]
        for title, breakdown in (("章节结构", paper.chapter_breakdown), ("难度结构", paper.difficulty_breakdown)):
            blocks.append(Paragraph(title, styles["Heading2"]))
            for row in breakdown:
                blocks.append(
                    Paragraph(
                        f"{self._escape(row.label)}：{row.question_count} 题，{self._score(row.score)} 分",
                        styles["Body"],
                    )
                )
        blocks.extend([PageBreak(), *self._pdf_sources(paper, styles)])
        return blocks

    def _pdf_sources(self, paper, styles) -> list:
        blocks = [
            Paragraph("内容来源与审核说明", styles["Heading"]),
            Paragraph(
                "本试卷保存题目版本快照；题目来源仅用于教师审计，PDF 原版式与原解析未进入导出内容。",
                styles["Notice"],
            ),
        ]
        for item in paper.items:
            blocks.append(
                Paragraph(
                    f"{item.position}. {self._escape(item.question.question_id)}｜"
                    f"{self._escape(item.question.source_document)}｜数学验证：通过｜教师审核："
                    f"{'通过' if item.question.review_status == 'approved' else '待审核'}",
                    styles["Small"],
                )
            )
        return blocks

    def _asset_path(self, paper_id: str, asset_id: str) -> Path:
        matches = list((self.asset_root / paper_id).glob(f"{asset_id}.*"))
        if len(matches) != 1:
            raise ExamPaperExportError(f"试卷图片快照缺失：{asset_id}")
        path = matches[0].resolve()
        if self.asset_root not in path.parents or not path.is_file():
            raise ExamPaperExportError(f"试卷图片路径无效：{asset_id}")
        return path

    def _configure_docx(self, document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.color.rgb = RGBColor.from_string(self.TEXT)
        normal.paragraph_format.line_spacing = 1.45
        normal.paragraph_format.space_after = Pt(5)

    def _docx_header_footer(self, section, edition: ExamPaperEdition) -> None:
        header = section.header.paragraphs[0]
        self._set_run(header.add_run(f"数研备课｜{self._edition_label(edition)}"), 8, self.MUTED)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_run(footer.add_run("第 "), 8, self.MUTED)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)
        self._set_run(footer.add_run(" 页"), 8, self.MUTED)

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), fill)
        properties.append(shade)

    @staticmethod
    def _set_run(run, size: float, color: str, *, bold: bool = False) -> None:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)

    def _resolve_pdf_fonts(self) -> tuple[Path, Path]:
        regular_candidates = [
            self.cjk_font_regular,
            Path("C:/Windows/Fonts/msyh.ttc"),
            Path("C:/Windows/Fonts/simsun.ttc"),
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        ]
        bold_candidates = [
            self.cjk_font_bold,
            Path("C:/Windows/Fonts/msyhbd.ttc"),
            Path("C:/Windows/Fonts/simhei.ttf"),
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"),
        ]
        regular = next((path for path in regular_candidates if path and path.is_file()), None)
        bold = next((path for path in bold_candidates if path and path.is_file()), None)
        if not regular:
            raise ExamPaperExportError("未找到可用中文字体，请配置 MATH_AI_CJK_FONT_REGULAR")
        return regular, bold or regular

    def _pdf_styles(self, regular: str, bold: str) -> dict[str, ParagraphStyle]:
        base = getSampleStyleSheet()
        return {
            "Title": ParagraphStyle("PaperTitle", parent=base["Title"], fontName=bold, fontSize=20, leading=29, textColor=colors.HexColor(f"#{self.TEXT}"), alignment=TA_CENTER, spaceAfter=7),
            "Meta": ParagraphStyle("PaperMeta", parent=base["BodyText"], fontName=regular, fontSize=9, leading=14, textColor=colors.HexColor(f"#{self.MUTED}"), alignment=TA_CENTER, spaceAfter=4),
            "Notice": ParagraphStyle("PaperNotice", parent=base["BodyText"], fontName=regular, fontSize=8.5, leading=14, textColor=colors.HexColor(f"#{self.MUTED}"), backColor=colors.HexColor("#F6F7FA"), borderPadding=7, spaceAfter=8),
            "Heading": ParagraphStyle("PaperHeading", parent=base["Heading2"], fontName=bold, fontSize=13, leading=19, textColor=colors.HexColor(f"#{self.BLUE}"), spaceBefore=12, spaceAfter=7, keepWithNext=True),
            "Heading2": ParagraphStyle("PaperHeading2", parent=base["Heading3"], fontName=bold, fontSize=10.5, leading=16, textColor=colors.HexColor(f"#{self.TEXT}"), spaceBefore=10, spaceAfter=4, keepWithNext=True),
            "Question": ParagraphStyle("PaperQuestion", parent=base["BodyText"], fontName=regular, fontSize=10.5, leading=18, textColor=colors.HexColor(f"#{self.TEXT}"), spaceAfter=5),
            "Option": ParagraphStyle("PaperOption", parent=base["BodyText"], fontName=regular, fontSize=9.5, leading=15, leftIndent=16, textColor=colors.HexColor(f"#{self.TEXT}"), spaceAfter=2),
            "Answer": ParagraphStyle("PaperAnswer", parent=base["BodyText"], fontName=regular, fontSize=9.5, leading=15, leftIndent=10, textColor=colors.HexColor("#176F60"), backColor=colors.HexColor(f"#{self.PALE_GREEN}"), borderPadding=6, spaceBefore=4, spaceAfter=4),
            "Solution": ParagraphStyle("PaperSolution", parent=base["BodyText"], fontName=regular, fontSize=9, leading=15, leftIndent=15, textColor=colors.HexColor(f"#{self.TEXT}"), spaceAfter=2),
            "Small": ParagraphStyle("PaperSmall", parent=base["BodyText"], fontName=regular, fontSize=7.5, leading=12, textColor=colors.HexColor(f"#{self.MUTED}"), spaceAfter=3),
            "Caption": ParagraphStyle("PaperCaption", parent=base["BodyText"], fontName=regular, fontSize=7.5, leading=11, textColor=colors.HexColor(f"#{self.MUTED}"), alignment=TA_CENTER, spaceAfter=5),
            "TableHead": ParagraphStyle("PaperTableHead", parent=base["BodyText"], fontName=bold, fontSize=7, leading=10, textColor=colors.HexColor(f"#{self.BLUE}"), alignment=TA_CENTER),
            "TableCell": ParagraphStyle("PaperTableCell", parent=base["BodyText"], fontName=regular, fontSize=6.5, leading=9, textColor=colors.HexColor(f"#{self.TEXT}")),
            "Body": ParagraphStyle("PaperBody", parent=base["BodyText"], fontName=regular, fontSize=9, leading=14, textColor=colors.HexColor(f"#{self.TEXT}"), spaceAfter=3),
        }

    @staticmethod
    def _escape(value: str) -> str:
        return html.escape(str(value)).replace("\n", "<br/>")

    @staticmethod
    def _score(value: float) -> str:
        return str(int(value)) if float(value).is_integer() else f"{value:g}"

    @staticmethod
    def _edition_label(edition: ExamPaperEdition) -> str:
        return {"student": "学生卷", "answer": "答案卷", "blueprint": "双向细目表"}[edition]

    @staticmethod
    def _question_type_label(question_type: str) -> str:
        return {
            "single_choice": "单选题",
            "multiple_choice": "多选题",
            "fill_blank": "填空题",
            "open_response": "解答题",
        }.get(question_type, "解答题")
